from __future__ import annotations

"""원본 판결문 + Easy-Read 합본 PDF 내보내기.

규칙:
- 원본이 .docx면 본문 구조를 유지해 지정 문구 뒤에 Easy-Read 문서를 삽입.
- 원본이 .pdf면 pdf2docx로 Word(.docx) 변환 후 동일 위치에 삽입.
- 최종 결과는 PDF로 반환.
"""

import io
import logging
import re
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docxcompose.composer import Composer
from pdf2docx import Converter

from backend.services.docx_to_pdf import convert_file_bytes_to_pdf

_MARKER_TEXT = "장애인 등을 위한 이해하기 쉬운(Easy-Read) 판결의 제공"
_REASON_HEADINGS = {"이유", "판결이유"}
logger = logging.getLogger(__name__)


def _convert_pdf_bytes_to_docx_with_pdf2docx(content: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        pdf_path = tmp_dir / "source.pdf"
        docx_path = tmp_dir / "converted.docx"
        pdf_path.write_bytes(content)

        converter = Converter(str(pdf_path))
        try:
            converter.convert(str(docx_path), start=0, end=None)
        finally:
            converter.close()

        if not docx_path.is_file():
            raise ValueError("pdf2docx 변환 결과 파일을 찾지 못했습니다.")
        return docx_path.read_bytes()


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text).replace("[", "").replace("]", "")


def _extract_element_text(element) -> str:
    parts: list[str] = []
    for node in element.iter():
        if node.tag == qn("w:t") and node.text:
            parts.append(node.text)
    return "".join(parts)


def _is_reason_heading(normalized_text: str) -> bool:
    if normalized_text in _REASON_HEADINGS:
        return True
    return normalized_text.startswith("이유") and len(normalized_text) <= 6


def _docx_bytes(doc: Document) -> bytes:
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _non_sect_children(doc: Document):
    return [child for child in doc._element.body.iterchildren() if child.tag != qn("w:sectPr")]


def _marker_insert_index(doc: Document) -> int:
    marker = _normalize_text(_MARKER_TEXT)
    children = _non_sect_children(doc)

    reason_idx: int | None = None
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = _normalize_text(_extract_element_text(child))
        if _is_reason_heading(text):
            reason_idx = idx
            break

    start_idx = reason_idx + 1 if reason_idx is not None else 0
    for idx in range(start_idx, len(children)):
        child = children[idx]
        if child.tag != qn("w:p"):
            continue
        text = _normalize_text(_extract_element_text(child))
        if marker in text:
            return idx + 1

    if reason_idx is not None:
        # New rule: if marker paragraph is missing after "이유", insert immediately after "이유".
        return reason_idx + 1

    return len(children)


def _slice_docx(base_docx: bytes, start: int, end: int) -> bytes:
    """Keep only body elements in [start:end], preserving relationships/media parts."""
    doc = Document(io.BytesIO(base_docx))
    body = doc._element.body
    children = [child for child in list(body.iterchildren()) if child.tag != qn("w:sectPr")]
    for idx, child in enumerate(children):
        if idx < start or idx >= end:
            body.remove(child)
    return _docx_bytes(doc)


def _compose_docx(parts: list[bytes]) -> bytes:
    master = Document(io.BytesIO(parts[0]))
    composer = Composer(master)
    for part in parts[1:]:
        composer.append(Document(io.BytesIO(part)))
    out = io.BytesIO()
    composer.save(out)
    return out.getvalue()


def _docx_with_inserted_easyread(base_docx: bytes, easyread_docx: bytes) -> bytes:
    base_doc = Document(io.BytesIO(base_docx))
    children = _non_sect_children(base_doc)
    insert_idx = _marker_insert_index(base_doc)

    # Preserve image/media relationships by composing complete documents,
    # not by deep-copying XML nodes between unrelated packages.
    # Insert immediately after marker/reason heading without adding page breaks.
    if insert_idx >= len(children):
        return _compose_docx([base_docx, easyread_docx])

    head = _slice_docx(base_docx, 0, insert_idx)
    tail = _slice_docx(base_docx, insert_idx, len(children))
    return _compose_docx([head, easyread_docx, tail])


def _build_docx_from_parsed_text(*, pages: list[str] | None, full_text: str | None) -> bytes:
    doc = Document()

    if pages:
        for i, page in enumerate(pages):
            lines = [line.strip() for line in page.splitlines() if line.strip()]
            for line in lines:
                doc.add_paragraph(line)
            if i < len(pages) - 1:
                pb = doc.add_paragraph()
                run = pb.add_run()
                run.add_break(WD_BREAK.PAGE)
    else:
        text = full_text or ""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        for line in lines:
            doc.add_paragraph(line)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def export_merged_docx(
    *,
    original_source_bytes: bytes,
    original_filename: str,
    easyread_docx_bytes: bytes,
    parsed_pages: list[str] | None = None,
    parsed_full_text: str | None = None,
) -> bytes:
    suffix = Path(original_filename).suffix.lower()

    if suffix == ".docx":
        base_docx = original_source_bytes
    elif suffix == ".pdf":
        try:
            base_docx = _convert_pdf_bytes_to_docx_with_pdf2docx(original_source_bytes)
        except Exception as exc:  # noqa: BLE001 - re-raise with user-facing context
            logger.error("pdf2docx PDF→DOCX 변환 실패: %s", exc)
            raise ValueError("PDF 원문을 pdf2docx로 Word 변환하지 못했습니다.") from exc
    else:
        # .doc 및 기타 포맷은 파싱 텍스트 기반 재구성으로 처리
        base_docx = _build_docx_from_parsed_text(
            pages=parsed_pages,
            full_text=parsed_full_text,
        )

    return _docx_with_inserted_easyread(base_docx, easyread_docx_bytes)


def export_merged_pdf(
    *,
    original_source_bytes: bytes,
    original_filename: str,
    easyread_docx_bytes: bytes,
    parsed_pages: list[str] | None = None,
    parsed_full_text: str | None = None,
) -> bytes:
    merged_docx = export_merged_docx(
        original_source_bytes=original_source_bytes,
        original_filename=original_filename,
        easyread_docx_bytes=easyread_docx_bytes,
        parsed_pages=parsed_pages,
        parsed_full_text=parsed_full_text,
    )
    return convert_file_bytes_to_pdf(merged_docx, "merged.docx")


def concat_pdf_bytes(first_pdf: bytes, second_pdf: bytes) -> bytes:
    """Append `second_pdf` pages after `first_pdf` and return merged PDF bytes."""
    out = fitz.open()
    first = fitz.open(stream=first_pdf, filetype="pdf")
    second = fitz.open(stream=second_pdf, filetype="pdf")
    try:
        out.insert_pdf(first)
        out.insert_pdf(second)
        return out.tobytes(garbage=3, deflate=True)
    finally:
        first.close()
        second.close()
        out.close()


def concat_many_pdf_bytes(parts: list[bytes]) -> bytes:
    """Concatenate multiple PDFs in order and return one PDF."""
    valid_parts = [p for p in parts if p]
    if not valid_parts:
        return b""
    if len(valid_parts) == 1:
        return valid_parts[0]

    out = fitz.open()
    try:
        for blob in valid_parts:
            src = fitz.open(stream=blob, filetype="pdf")
            try:
                out.insert_pdf(src)
            finally:
                src.close()
        return out.tobytes(garbage=3, deflate=True)
    finally:
        out.close()


def split_parsed_text_for_easyread_insert(text: str) -> tuple[str, str]:
    """Split parsed original text into (before_insert, after_insert)."""
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if not lines:
        return "", ""

    marker = _normalize_text(_MARKER_TEXT)
    reason_idx: int | None = None
    for idx, line in enumerate(lines):
        if _is_reason_heading(_normalize_text(line)):
            reason_idx = idx
            break

    start_idx = reason_idx + 1 if reason_idx is not None else 0
    insert_idx = len(lines)
    for idx in range(start_idx, len(lines)):
        if marker in _normalize_text(lines[idx]):
            insert_idx = idx + 1
            break

    if reason_idx is not None and insert_idx == len(lines):
        insert_idx = reason_idx + 1

    before = "\n".join(lines[:insert_idx]).strip()
    after = "\n".join(lines[insert_idx:]).strip()
    return before, after


def find_insert_page_index(parsed_pages: list[str]) -> int:
    """Find page index after which Easy-Read pages should be inserted."""
    if not parsed_pages:
        return -1

    marker = _normalize_text(_MARKER_TEXT)
    reason_page: int | None = None

    for idx, page in enumerate(parsed_pages):
        for line in page.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
            if _is_reason_heading(_normalize_text(line)):
                reason_page = idx
                break
        if reason_page is not None:
            break

    start_idx = reason_page if reason_page is not None else 0
    for idx in range(start_idx, len(parsed_pages)):
        if marker in _normalize_text(parsed_pages[idx]):
            return idx

    if reason_page is not None:
        return reason_page

    return len(parsed_pages) - 1


def insert_pdf_after_page(original_pdf: bytes, insert_pdf: bytes, after_page_index: int) -> bytes:
    """Insert `insert_pdf` after `after_page_index` of `original_pdf` and return merged bytes."""
    source = fitz.open(stream=original_pdf, filetype="pdf")
    extra = fitz.open(stream=insert_pdf, filetype="pdf")
    out = fitz.open()
    try:
        total = source.page_count
        if total == 0:
            out.insert_pdf(extra)
            return out.tobytes(garbage=3, deflate=True)

        if after_page_index < 0:
            out.insert_pdf(extra)
            out.insert_pdf(source)
            return out.tobytes(garbage=3, deflate=True)

        if after_page_index >= total - 1:
            out.insert_pdf(source)
            out.insert_pdf(extra)
            return out.tobytes(garbage=3, deflate=True)

        out.insert_pdf(source, from_page=0, to_page=after_page_index)
        out.insert_pdf(extra)
        out.insert_pdf(source, from_page=after_page_index + 1, to_page=total - 1)
        return out.tobytes(garbage=3, deflate=True)
    finally:
        source.close()
        extra.close()
        out.close()
